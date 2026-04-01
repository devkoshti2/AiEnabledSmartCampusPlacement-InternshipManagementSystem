"""
Resume Parser - Complete Version v2
====================================
Handles ANY resume format:
- Table-based / two-column DOCX resumes
- Single-column simple resumes
- PDFs (text-based)
- Any domain: CS, Mechanical, Civil, Electrical, MBA, etc.

Key improvements over v1:
1. Domain-specific skills DB: Mechanical, Civil, Electrical, MBA, etc.
2. Smart fallback: extracts explicitly listed skills from Skills section
   even if not in DB (catches domain-specific / custom skills)
3. Hyphen-space normalization: 'problem-solving' = 'problem solving'
4. Better name extractor: handles single-word names, skips institution names
5. Fixed CGPA regex: handles 8.75, 7.25, 9.2 correctly
6. No false positives: 'c', 'r', 'go', 'ai' don't match everything
7. git + github both extracted correctly
8. java + javascript both extracted correctly
"""

try:
    from pypdf import PdfReader
except ImportError:
    try:
        import PyPDF2
        PdfReader = PyPDF2.PdfReader
    except ImportError:
        PdfReader = None

try:
    import docx
except ImportError:
    docx = None

import re
import os
import logging
from typing import Dict, List, Any

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeParser:
    """Universal Resume Parser"""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        if PdfReader is None:
            logger.error("❌ No PDF library. Install: pip install pypdf")
            return ""
        text = ""
        try:
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            logger.info(f"✅ PDF extracted: {len(text)} chars")
            return text
        except Exception as e:
            logger.error(f"❌ PDF error: {e}")
            return ""

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """DOCX extraction — paragraphs + tables, no duplicate runs"""
        if docx is None:
            logger.error("❌ python-docx not installed.")
            return ""
        text = ""
        try:
            doc = docx.Document(file_path)
            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # Tables (important for 2-column resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text += " | ".join(row_texts) + "\n"
            logger.info(f"✅ DOCX extracted: {len(text)} chars")
            if text:
                logger.info(f"📄 Preview: {text[:200]}")
            return text
        except Exception as e:
            logger.error(f"❌ DOCX error: {e}")
            return ""

    @staticmethod
    def extract_email(text: str) -> str:
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', text)
        return emails[0] if emails else ""

    @staticmethod
    def extract_phone(text: str) -> str:
        patterns = [
            r'\b\+?91[\s\-]?\d{10}\b',
            r'\b\d{10}\b',
            r'\b\d{3}[.\-\s]\d{3}[.\-\s]\d{4}\b',
            r'\+\d{1,3}[\s\-]\d{3,5}[\s\-]\d{3,5}[\s\-]?\d{3,5}',
        ]
        for p in patterns:
            m = re.findall(p, text)
            if m:
                return m[0].strip()
        return ""

    @staticmethod
    def extract_name(text: str) -> str:
        """Extract name from first 15 lines"""
        skip_words = [
            'resume', 'cv', 'curriculum', 'vitae', 'email', 'phone',
            'address', 'linkedin', 'github', '@', 'http', 'www',
            'education', 'experience', 'skills', 'summary', 'objective',
            'page', 'document', 'created', 'updated', 'contact', 'profile',
            'college', 'university', 'institute', 'school', 'academy',
            'engineer', 'developer', 'software', 'technology', 'technologies',
            'enthusiast', 'professional', 'specialist', 'consultant',
            'analyst', 'intern', 'manager', 'lead', 'senior', 'junior',
            'mechanical', 'electrical', 'civil', 'computer', 'information',
        ]
        for line in text.split('\n')[:15]:
            line = line.strip()
            if not line or len(line) < 2 or len(line) > 50:
                continue
            if any(w in line.lower() for w in skip_words):
                continue
            if re.match(r'^[A-Za-z\s.\-]+$', line):
                words = line.split()
                if 1 <= len(words) <= 4:
                    if all(w[0].isupper() for w in words if w):
                        logger.info(f"✅ Name: {line}")
                        return line
        return ""


class SkillExtractor:
    """
    Universal Skill Extractor v2

    Strategy:
    1. Match against comprehensive skills DB (all domains)
    2. Smart fallback: extract explicitly listed skills from Skills section
       (handles domain-specific skills not in DB)
    3. Merge both, deduplicate
    """

    # Skills requiring strict word-boundary (avoid substring false positives)
    STRICT_BOUNDARY_SKILLS = {
        'c', 'r', 'ai', 'c#', 'c++', 'java', 'sql', 'php',
        'css', 'xml', 'ios', 'aws', 'git', 'api', 'nlp', 'rag',
        'gpt', 'llm', 'fea', 'cad', 'plc', 'erp', 'crm', 'bi',
        'hr', 'pr', 'assembly',
    }

    # Pairs that are BOTH valid even though one is substring of other
    KEEP_PAIRS = {
        # Version control
        ('git', 'github'), ('git', 'gitlab'), ('git', 'bitbucket'),
        # JS ecosystem
        ('node', 'node.js'), ('react', 'react native'), ('react', 'next.js'),
        ('spring', 'spring boot'), ('vue', 'vue.js'), ('java', 'javascript'),
        # CAD
        ('cad', 'autocad'), ('cad', 'cad drafting'),
        # C language - keep 'c' alongside c++, embedded c, etc.
        ('c', 'c++'), ('c', 'c#'), ('c', 'embedded c'), ('c', 'embedded c basics'),
        ('c', 'embedded c++'), ('c', 'objective-c'),
        # R language
        ('r', 'react'), ('r', 'rust'),
        # Embedded
        ('embedded c', 'embedded c basics'), ('embedded c', 'embedded c++'),
        ('embedded c', 'embedded systems'),
        ('microcontroller', 'microcontroller programming'),
        # Signal processing
        ('signal processing', 'digital signal processing'),
        ('signal processing', 'signal processing basics'),
        # Machine learning
        ('machine learning', 'deep learning'), ('neural networks', 'convolutional neural network'),
        # Data
        ('data', 'data structures'), ('data', 'data science'), ('data', 'data mining'),
        ('data structures', 'data structures and algorithms'),
        # Python
        ('python', 'python programming'),
        # Communication
        ('communication', 'communication skills'),
        ('programming', 'microcontroller programming'),
    }

    def __init__(self):
        self.skills_db = self._build_skills_db()
        self.common_words = self._build_common_words()
        self._patterns = self._compile_patterns()

    def _build_common_words(self) -> set:
        return {
            'and', 'the', 'for', 'with', 'using', 'etc', 'skills', 'skill',
            'education', 'experience', 'project', 'projects', 'work', 'job',
            'company', 'position', 'role', 'title', 'date', 'year', 'month',
            'description', 'responsibilities', 'achievements', 'summary',
            'objective', 'profile', 'contact', 'reference',
            'certification', 'certificate', 'training',
            'university', 'college', 'school', 'degree', 'bachelor', 'master',
            'phd', 'diploma', 'cgpa', 'gpa', 'percentage',
            'marks', 'grade', 'score', 'result', 'passing',
            'currently', 'present', 'till', 'until', 'from', 'to',
            'manager', 'lead', 'senior', 'junior', 'trainee', 'intern',
            'full-time', 'part-time', 'contract', 'freelance', 'remote',
            'placement', 'beginner',
        }

    def _build_skills_db(self) -> List[str]:
        """
        Comprehensive multi-domain skills database.
        Sorted longest-first so specific skills match before short ones.
        """
        skills = [
            # ── Programming Languages ──────────────────────────────────
            'python', 'javascript', 'typescript', 'java', 'c++', 'c#', 'c',
            'ruby', 'php', 'swift', 'kotlin', 'golang', 'go language', 'rust', 'scala', 'perl',
            'r', 'matlab', 'dart', 'elixir', 'haskell', 'clojure',
            'objective-c', 'vb.net', 'groovy', 'lua', 'julia',
            'assembly language', 'x86 assembly', 'fortran', 'cobol',

            # ── Web Frontend ───────────────────────────────────────────
            'html', 'html5', 'css', 'css3', 'react', 'angular', 'vue.js',
            'next.js', 'nuxt.js', 'svelte', 'gatsby', 'htmx',
            'jquery', 'bootstrap', 'tailwind css', 'tailwind',
            'sass', 'scss', 'webpack', 'babel', 'redux', 'pwa', 'webassembly',

            # ── Web Backend ────────────────────────────────────────────
            'node.js', 'express.js', 'express', 'django', 'flask',
            'spring boot', 'spring', 'asp.net', 'fastapi', 'laravel',
            'rest api', 'restful api', 'graphql', 'soap', 'ajax',
            'microservices', 'websockets',

            # ── Databases ──────────────────────────────────────────────
            'mysql', 'postgresql', 'mongodb', 'sqlite', 'redis',
            'oracle', 'cassandra', 'elasticsearch', 'firebase',
            'mariadb', 'dynamodb', 'neo4j', 'influxdb',
            'sql server', 'snowflake', 'bigquery', 'redshift',
            'cosmos db', 'sql',

            # ── Cloud & DevOps ─────────────────────────────────────────
            'aws', 'azure', 'gcp', 'google cloud',
            'docker', 'kubernetes', 'jenkins', 'git', 'github',
            'gitlab', 'bitbucket', 'ansible', 'terraform',
            'prometheus', 'grafana', 'ci/cd', 'devops',
            'linux', 'ubuntu', 'windows', 'macos', 'unix',

            # ── Data Science / AI / ML ─────────────────────────────────
            'machine learning', 'deep learning', 'data science',
            'artificial intelligence', 'computer vision',
            'natural language processing', 'nlp',
            'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'opencv',
            'pandas', 'numpy', 'matplotlib', 'seaborn', 'plotly',
            'langchain', 'spacy', 'nltk', 'transformers', 'hugging face',
            'bert', 'gpt', 'llama', 'rag', 'llm', 'genai', 'ai',
            'data mining', 'statistics', 'neural networks',
            'reinforcement learning', 'transfer learning',
            'data augmentation', 'convolutional neural network',
            'cnn', 'rnn', 'lstm',

            # ── Mobile ─────────────────────────────────────────────────
            'android', 'ios', 'react native', 'flutter', 'xamarin',
            'swiftui', 'jetpack compose', 'ionic', 'cordova',
            'android studio', 'xcode',

            # ── Testing ────────────────────────────────────────────────
            'selenium', 'pytest', 'junit', 'jest', 'cypress',
            'playwright', 'postman', 'jmeter',
            'unit testing', 'integration testing', 'tdd', 'bdd', 'qa',

            # ── Tools & IDEs ───────────────────────────────────────────
            'vs code', 'visual studio', 'intellij', 'eclipse',
            'jira', 'confluence', 'trello', 'asana', 'notion',
            'figma', 'adobe xd', 'canva', 'photoshop', 'sketch',
            'microsoft office', 'excel', 'powerpoint', 'word',
            'version control',

            # ── Security & Networking ──────────────────────────────────
            'cybersecurity', 'ethical hacking', 'penetration testing',
            'networking', 'tcp/ip', 'dns', 'ssl/tls', 'oauth', 'jwt',
            'cryptography', 'firewall', 'vpn',
            'wireshark', 'nmap', 'metasploit', 'burp suite',

            # ── Methodologies ──────────────────────────────────────────
            'agile', 'scrum', 'kanban', 'waterfall',
            'object oriented programming', 'oop',
            'design patterns', 'system design',
            'data structures', 'algorithms',

            # ── Soft Skills ────────────────────────────────────────────
            'communication', 'leadership', 'teamwork', 'team collaboration',
            'problem solving', 'critical thinking', 'time management',
            'presentation skills', 'public speaking', 'analytical thinking',
            'decision making', 'adaptability', 'creativity', 'collaboration',
            'emotional intelligence', 'mentoring', 'negotiation',
            'logical thinking', 'quick learner', 'self-motivated',
            'interpersonal skills', 'work ethic', 'conflict resolution',

            # ══════════════════════════════════════════════════════════
            # ── MECHANICAL ENGINEERING ────────────────────────────────
            # ══════════════════════════════════════════════════════════
            'autocad', 'solidworks', 'ansys', 'catia', 'creo',
            'unigraphics', 'nx', 'inventor', 'fusion 360',
            'pro/engineer', 'hypermesh', 'abaqus', 'nastran',
            'finite element analysis', 'fea', 'fem',
            'computational fluid dynamics', 'cfd',
            'thermodynamics', 'heat transfer', 'fluid mechanics',
            'fluid dynamics', 'fluid flow',
            'heat transfer & fluid flow', 'heat exchanger design',
            'thermal analysis', 'stress analysis',
            'thermal simulation', 'stress simulation',
            'thermal & stress simulation',
            'structural analysis', 'structural optimization',
            'fatigue analysis', 'vibration analysis',
            'modal analysis', 'buckling analysis',
            'static analysis', 'dynamic analysis',
            'manufacturing processes', 'manufacturing',
            'cnc machining', 'cnc programming', 'cnc',
            'gd&t', 'geometric dimensioning',
            'cad drafting', 'technical drawings',
            'cad drafting & technical drawings',
            '3d modeling', 'parametric design', 'parametric modeling',
            '3d modeling & parametric design',
            'sheet metal design', 'welding', 'casting',
            'material science', 'material selection',
            'material selection & optimization',
            'strength of materials', 'mechanics of materials',
            'machine design', 'machine elements',
            'gearbox design', 'bearing design', 'shaft design',
            'product design', 'product development',
            'rapid prototyping', '3d printing', 'additive manufacturing',
            'reverse engineering', 'tolerance analysis',
            'quality control', 'lean manufacturing', 'six sigma',
            'kaizen', 'value engineering',
            'robotics', 'automation', 'mechatronics',
            'plc programming', 'plc', 'scada',
            'hydraulics', 'pneumatics',
            'renewable energy', 'solar energy', 'wind energy',
            'hvac', 'refrigeration',
            'ic engine', 'automobile engineering',
            'piping design', 'pressure vessel design',

            # ══════════════════════════════════════════════════════════
            # ── CIVIL ENGINEERING ─────────────────────────────────────
            # ══════════════════════════════════════════════════════════
            'staad pro', 'etabs', 'sap2000', 'revit', 'navisworks',
            'primavera', 'ms project',
            'structural design', 'rcc design', 'steel design',
            'construction management', 'project management',
            'quantity surveying', 'estimation',
            'surveying', 'total station', 'gps surveying',
            'geotechnical engineering', 'soil mechanics',
            'foundation design', 'retaining wall design',
            'traffic engineering', 'transportation engineering',
            'environmental engineering', 'water treatment',
            'sewage treatment', 'wastewater treatment',
            'autocad civil', 'civil 3d',
            'concrete mix design', 'mix design',

            # ══════════════════════════════════════════════════════════
            # ── ELECTRICAL / ELECTRONICS ENGINEERING ─────────────────
            # ══════════════════════════════════════════════════════════
            'matlab simulink', 'simulink', 'labview',
            'multisim', 'proteus', 'orcad', 'altium designer',
            'eagle', 'kicad',
            'vhdl', 'verilog', 'fpga', 'embedded systems',
            'microcontroller programming', 'microcontroller', 'arduino', 'raspberry pi',
            'arm cortex', 'stm32', 'avr', 'pic microcontroller',
            'circuit design', 'pcb design', 'pcb layout',
            'power electronics', 'power systems',
            'electric drives', 'motor control',
            'signal processing', 'digital signal processing', 'dsp',
            'control systems', 'pid control',
            'iot', 'internet of things',
            'rtos', 'embedded c basics', 'embedded c', 'embedded c++',
            'bluetooth', 'wi-fi', 'zigbee', 'lora', 'mqtt',
            'sensor fusion', 'robotics process automation',
            'high voltage engineering', 'transformer design',

            # ══════════════════════════════════════════════════════════
            # ── MBA / BUSINESS / FINANCE ──────────────────────────────
            # ══════════════════════════════════════════════════════════
            'financial analysis', 'financial modelling', 'financial modeling',
            'accounting', 'tally', 'tally erp', 'sap',
            'erp', 'crm', 'salesforce', 'sap fico', 'sap mm',
            'business analysis', 'business intelligence', 'bi',
            'power bi', 'tableau', 'qlik', 'looker',
            'market research', 'marketing strategy',
            'digital marketing', 'seo', 'sem', 'social media marketing',
            'content marketing', 'google analytics',
            'supply chain management', 'logistics', 'inventory management',
            'operations management', 'strategic planning',
            'business development', 'sales', 'customer relationship',
            'human resources', 'recruitment', 'talent acquisition',
            'performance management', 'payroll',
            'investment banking', 'equity research',
            'portfolio management', 'risk management',
            'derivatives', 'valuation', 'mergers and acquisitions',

            # ══════════════════════════════════════════════════════════
            # ── CHEMICAL ENGINEERING ──────────────────────────────────
            # ══════════════════════════════════════════════════════════
            'aspen plus', 'aspen hysys', 'hysys', 'chemcad',
            'process simulation', 'process design',
            'mass transfer', 'reaction engineering',
            'process control', 'distillation', 'extraction',
            'polymer science', 'petrochemicals',

            # ══════════════════════════════════════════════════════════
            # ── GENERAL / CROSS-DOMAIN ────────────────────────────────
            # ══════════════════════════════════════════════════════════
            'ms office', 'microsoft office', 'word', 'excel', 'powerpoint',
            'latex', 'technical writing', 'report writing',
            'research methodology', 'literature review',
            'patent filing', 'intellectual property',
        ]
        # Sort longest first — prevents short skills from blocking longer ones
        return sorted(list(dict.fromkeys(skills)), key=len, reverse=True)

    def _compile_patterns(self) -> Dict[str, re.Pattern]:
        """Compile regex patterns — strict boundary for ambiguous short skills"""
        patterns = {}
        for skill in self.skills_db:
            sl = skill.lower()
            if sl in self.STRICT_BOUNDARY_SKILLS:
                escaped = re.escape(sl)
                pat = r'(?<![a-z0-9])' + escaped + r'(?![a-z0-9])'
            elif ' ' in sl or '&' in sl:
                # Multi-word: allow hyphen or space between words
                parts = [re.escape(w) for w in re.split(r'[\s&]+', sl) if w]
                pat = r'\b' + r'[\s\-&]+'.join(parts) + r'\b'
            else:
                escaped = re.escape(sl)
                pat = r'\b' + escaped + r'\b'
            try:
                patterns[sl] = re.compile(pat, re.IGNORECASE)
            except re.error:
                patterns[sl] = re.compile(re.escape(sl), re.IGNORECASE)
        return patterns

    def _extract_skills_section_raw(self, text: str) -> List[str]:
        """
        Extract explicitly listed skills from Skills section ONLY.

        Key design:
        1. Stop immediately at Project/Experience content (dates, long lines)
        2. Table format "Last Skill | Project" -> take before pipe, then stop
        3. Only accept short items (<=60 chars), no sentence structure
        4. Reject items that look like project description text
        """
        raw_items = []

        # ── STEP 1: Find and extract only the Skills section lines ─────
        lines = text.split('\n')
        in_skills = False
        skills_lines = []

        STOP_HEADINGS = re.compile(
            r'^(?:education|experience|projects?|certifications?|objective|'
            r'summary|contact|achievements?|awards?|languages?|references?|'
            r'internship|publications?|research|hobbies|interests?)',
            re.IGNORECASE
        )
        DATE_LINE = re.compile(
            r'\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b.{0,10}\d{4}'
            r'|\b\d{4}\s*[-–]|\b20\d{2}\b.{0,5}\b20\d{2}\b',
            re.IGNORECASE
        )
        LONG_SENTENCE = re.compile(r'.{80,}')

        for line in lines:
            ls = line.strip()
            # Detect Skills heading (standalone "Skills" or "Technical Skills")
            if re.match(r'^(?:technical\s+)?skills?\s*[:\|]?\s*$', ls, re.IGNORECASE):
                in_skills = True
                continue

            if in_skills:
                if not ls:
                    continue
                # Stop at section headings
                if STOP_HEADINGS.match(ls):
                    break
                # Stop at date patterns (project/experience section started)
                if DATE_LINE.search(ls):
                    break
                # Stop at long description lines
                if LONG_SENTENCE.match(ls):
                    break
                # Table pipe: "Last Skill | Project" -> take before |, stop
                if '|' in ls:
                    before_pipe = ls.split('|')[0].strip()
                    after_pipe  = ls.split('|', 1)[1].strip()
                    if STOP_HEADINGS.match(after_pipe):
                        if before_pipe:
                            skills_lines.append(before_pipe)
                        break  # section ends here
                    else:
                        if before_pipe:
                            skills_lines.append(before_pipe)
                        continue
                skills_lines.append(ls)

        # ── STEP 2: Parse each line into individual skill tokens ────────
        def parse_skill_line(item: str) -> List[str]:
            results = []
            item = re.sub(r'^[\-\*•]+\s*', '', item).strip()
            item = re.sub(r'^\d+[\.)]+\s+', '', item).strip()
            # Split comma-separated items on a single line
            parts = re.split(r',\s*', item)
            for part in parts:
                part = part.strip()
                if not part or len(part) > 60:
                    continue
                # Version with parenthetical content removed
                clean = re.sub(r'\s*\(.*?\)', '', part).strip()
                if clean and clean != part:
                    results.append(clean)    # e.g. "Database Basics"
                if part:
                    results.append(part)     # e.g. "Database Basics (MySQL/PostgreSQL)"
            return results

        for line in skills_lines:
            for item in parse_skill_line(line):
                if item and 1 <= len(item) <= 60:
                    raw_items.append(item)

        # ── STEP 3: Inline "Skills: X, Y, Z" format fallback ───────────
        inline = re.findall(
            r'(?:technical\s+skills?|skills?|expertise|technologies?)'
            r'\s*:\s*([^\n]{5,200})',
            text, re.IGNORECASE
        )
        for line in inline:
            for item in re.split(r'[,;]+', line):
                item = item.strip()
                if item and 1 <= len(item) <= 60:
                    raw_items.append(item)

        # ── STEP 4: Filter noise ────────────────────────────────────────
        NOISE_WORDS = re.compile(
            r'\b(?:project|capstone|academic|personal|internship|company|pvt|ltd|'
            r'jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|'
            r'responsible|developed|implemented|designed|created|built|'
            r'bachelor|master|b\.tech|m\.tech|cgpa|gpa|college|university)\b',
            re.IGNORECASE
        )
        # Reject items that contain project-description phrases
        PROJECT_PHRASES = re.compile(
            r'records?\.|features like|adding books|calculating|returning books|'
            r'searching by|issuing|hosting|focusing on|clean design|fast loading|'
            r'for backend|for frontend|for styling|for structure|responsiveness\)|'
            r'crud operations|user login\b|admin dashboard|professional presentation',
            re.IGNORECASE
        )

        # ── STEP 5: Deduplicate and return ──────────────────────────────
        seen = set()
        result = []
        for item in raw_items:
            key = item.lower().strip()
            if key in seen or key in self.common_words:
                continue
            if NOISE_WORDS.search(key):
                continue
            if PROJECT_PHRASES.search(key):
                continue
            # Reject if too long and contains verb phrases (description text)
            if len(key) > 40 and re.search(r'\b(?:using|with|for|and|the|to|of)\b', key):
                continue
            seen.add(key)
            result.append(item.strip())

        logger.info(f"📋 Raw section skills found: {len(result)} — {result[:10]}")
        return result


    def _normalize(self, text: str) -> str:
        """Normalize text for matching: collapse whitespace, handle hyphens"""
        text = text.lower().strip()
        # Normalize multiple spaces/newlines
        text = re.sub(r'\s+', ' ', text)
        return text

    def extract_skills(self, text: str) -> List[str]:
        """
        Main skill extraction — 3-step approach:

        Step 1: DB matching via regex on full text
        Step 2: Extract raw items from Skills section → match against DB
        Step 3: Add unmatched section items as-is (catches domain-specific skills)
        Final:  Deduplicate, remove substrings, return clean sorted list
        """
        if not text:
            logger.warning("⚠️ Empty text")
            return []

        logger.info(f"📝 Extracting skills (text length: {len(text)})")
        text_norm = self._normalize(text)
        # Also create a version with hyphens replaced by spaces for matching
        text_dehyphen = text_norm.replace('-', ' ')

        found_from_db = []
        matched_lower = set()

        # ── Step 1: DB regex matching ──────────────────────────────────
        for skill in self.skills_db:
            sl = skill.lower()
            if sl in matched_lower:
                continue
            pat = self._patterns.get(sl)
            if pat and (pat.search(text_norm) or pat.search(text_dehyphen)):
                found_from_db.append(sl)
                matched_lower.add(sl)

        # ── Step 2: Skills section → DB match ─────────────────────────
        section_raw = self._extract_skills_section_raw(text)
        for raw in section_raw:
            raw_norm = self._normalize(raw)
            raw_dehyphen = raw_norm.replace('-', ' ')

            if raw_norm in matched_lower:
                continue

            # Try exact match first
            matched = False
            for skill in self.skills_db:
                sl = skill.lower()
                if sl in matched_lower:
                    continue
                if sl == raw_norm or sl == raw_dehyphen:
                    found_from_db.append(sl)
                    matched_lower.add(sl)
                    matched = True
                    break

            if matched:
                continue

            # Try DB regex against this single raw item
            for skill in self.skills_db:
                sl = skill.lower()
                if sl in matched_lower:
                    continue
                pat = self._patterns.get(sl)
                if pat and (pat.search(raw_norm) or pat.search(raw_dehyphen)):
                    found_from_db.append(sl)
                    matched_lower.add(sl)
                    matched = True
                    break

        # ── Step 3: Add unmatched section items as-is ─────────────────
        # This is the "any format" safety net — if a skill is in the resume
        # Skills section but not in our DB, we still include it
        # Words that indicate noise / non-skill lines
        NOISE_PATTERNS = [
            r'project', r'capstone', r'academic', r'personal', r'jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec',
            r'\d{4}', r'internship', r'company', r'limited', r'pvt', r'ltd',
            r'responsible', r'developed', r'implemented', r'designed', r'created',
            r'bachelor|master|b\.tech|m\.tech|cgpa|gpa|college|university',
        ]
        noise_re = re.compile('|'.join(NOISE_PATTERNS), re.IGNORECASE)

        extra_skills = []
        for raw in section_raw:
            raw_norm = self._normalize(raw)
            raw_dehyphen = raw_norm.replace('-', ' ')

            # Skip noise
            if noise_re.search(raw_norm):
                continue

            # Skip if already matched to DB
            already_covered = False
            for matched_sl in matched_lower:
                if (raw_norm in matched_sl or matched_sl in raw_norm or
                        raw_dehyphen in matched_sl or matched_sl in raw_dehyphen):
                    already_covered = True
                    break

            if not already_covered and raw_norm not in self.common_words:
                if len(raw_norm) >= 1:
                    extra_skills.append(raw_norm)
                    matched_lower.add(raw_norm)

        all_found = found_from_db + extra_skills

        # ── Final: Remove redundant substrings ────────────────────────
        # Only remove a skill if it appears as a COMPLETE WORD inside another skill
        # e.g. 'machine' inside 'machine learning' — yes, remove
        # But 'c' inside 'communication' — NO, 'c' is not a whole word there
        found_set = set(all_found)
        final = []
        for skill in all_found:
            is_redundant = False
            for other in found_set:
                if skill == other:
                    continue
                # Check: is 'skill' a whole-word match inside 'other'?
                escaped = re.escape(skill)
                word_in_other = re.search(r'(?<![a-z0-9])' + escaped + r'(?![a-z0-9])', other)
                if word_in_other and other in found_set:
                    pair = tuple(sorted([skill, other]))
                    if pair not in self.KEEP_PAIRS and (skill, other) not in self.KEEP_PAIRS:
                        is_redundant = True
                        break
            if not is_redundant and skill not in self.common_words:
                final.append(skill)

        # Deduplicate preserving order, then sort
        seen = set()
        deduped = []
        for s in final:
            if s not in seen:
                seen.add(s)
                deduped.append(s)
        deduped.sort()

        logger.info(f"✅ Total skills extracted: {len(deduped)}")
        if deduped:
            logger.info(f"📋 Skills: {', '.join(deduped[:25])}")
        return deduped

    def extract_education(self, text: str) -> Dict[str, str]:
        """Education extraction with improved CGPA and semester parsing"""
        education = {
            'degree': '', 'branch': '', 'institution': '',
            'year': '', 'cgpa': '', 'semester': ''
        }
        text_lower = text.lower()

        # Degree
        degree_map = {
            'BTECH': r'b\.?tech|bachelor\s+of\s+technology',
            'MTECH': r'm\.?tech|master\s+of\s+technology',
            'BE':    r'\bb\.?e\.?\b|bachelor\s+of\s+engineering',
            'ME':    r'\bm\.?e\.?\b|master\s+of\s+engineering',
            'BSC':   r'b\.?sc|bachelor\s+of\s+science',
            'MSC':   r'm\.?sc|master\s+of\s+science',
            'MCA':   r'\bmca\b|master\s+of\s+computer\s+applications',
            'BCA':   r'\bbca\b|bachelor\s+of\s+computer\s+applications',
            'PHD':   r'ph\.?d|doctor\s+of\s+philosophy',
            'MBA':   r'\bmba\b|master\s+of\s+business',
            'BBA':   r'\bbba\b|bachelor\s+of\s+business',
            'BCOM':  r'\bb\.?com\b|bachelor\s+of\s+commerce',
            'MCOM':  r'\bm\.?com\b|master\s+of\s+commerce',
        }
        for degree, pat in degree_map.items():
            if re.search(pat, text_lower):
                education['degree'] = degree
                break

        # Branch
        branch_map = {
            'CSE':    r'computer\s*science|information\s*technology|\b(?:cs|cse|it)\b',
            'ECE':    r'electronics\s*(?:and\s*)?communication|\bece\b',
            'EEE':    r'electrical\s*(?:and\s*)?electronics|\beee\b',
            'MECH':   r'mechanical|\bmech\b',
            'CIVIL':  r'\bcivil\b',
            'CHEM':   r'chemical\s+engineering|\bchem\b',
            'AIDS':   r'artificial\s+intelligence\s+(?:and\s+)?data\s+science',
            'AI':     r'artificial\s+intelligence|\bai\s*&\s*ml\b',
            'DATA':   r'data\s+science',
            'CYBER':  r'cyber\s*security',
            'MBA':    r'business\s+administration|management',
            'BCOM':   r'commerce|\baccounting\b',
        }
        for branch, pat in branch_map.items():
            if re.search(pat, text_lower):
                education['branch'] = branch
                break

        # CGPA
        cgpa_pats = [
            r'(?:cgpa|gpa|cpi)\s*[:\-]?\s*(\d+\.\d{1,2})',
            r'(\d+\.\d{1,2})\s*/\s*10',
            r'(\d+\.\d{1,2})\s*(?:cgpa|gpa|out\s+of\s+10)',
            r'cgpa\s+(\d+\.\d{1,2})',
        ]
        for pat in cgpa_pats:
            m = re.search(pat, text_lower)
            if m:
                val = float(m.group(1))
                if 0 < val <= 10:
                    education['cgpa'] = str(val)
                    break

        # Semester
        sem_pats = [
            r'semester\s*[:\-]?\s*(\d{1,2})',
            r'(\d{1,2})\s*(?:st|nd|rd|th)\s*semester',
            r'sem\s*[:\-]?\s*(\d{1,2})',
        ]
        for pat in sem_pats:
            m = re.search(pat, text_lower)
            if m:
                sem = int(m.group(1))
                if 1 <= sem <= 12:
                    education['semester'] = str(sem)
                    break

        # Year
        years = re.findall(r'\b(20\d{2})\b', text)
        if years:
            education['year'] = max(years)

        return education

    def extract_experience(self, text: str) -> List[Dict[str, str]]:
        experiences = []
        exp_match = re.search(
            r'(?:experience|work\s+experience|employment|internship)\s*[:\|]?\s*\n(.*?)(?:\n\n\S|\Z)',
            text, re.IGNORECASE | re.DOTALL
        )
        if exp_match:
            for entry in re.split(r'\n\s*[-•*]\s*|\n\s*\d+\.\s*', exp_match.group(1)):
                entry = entry.strip()
                if len(entry) > 20:
                    lines = entry.split('\n')
                    experiences.append({
                        'title': lines[0][:80],
                        'description': entry[:300]
                    })
        return experiences

    def extract_projects(self, text: str) -> List[Dict[str, str]]:
        projects = []
        proj_match = re.search(
            r'(?:projects?|academic\s+projects?|personal\s+projects?)\s*[:\|]?\s*\n(.*?)(?:\n\n\S|\Z)',
            text, re.IGNORECASE | re.DOTALL
        )
        if proj_match:
            for entry in re.split(r'\n\s*[-•*]\s*|\n\s*\d+\.\s*', proj_match.group(1)):
                entry = entry.strip()
                if len(entry) < 10:
                    continue
                techs = self.extract_skills(entry)
                lines = entry.split('\n')
                projects.append({
                    'name': lines[0][:60],
                    'description': entry[:300],
                    'technologies': ', '.join(techs[:6]) if techs else ''
                })
        return projects